from PyPDF2 import PdfReader
import spacy
from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from docx import Document

root = Tk()
vetor_palavrasChave = []


class Funcs():
    def __init__(self):
        self.caminho_arquivo = ""
        self.textoCompleto = ""
        self.nlp = None
        self.doc = None
        self.palavras_chaves = []
        self.encontrados = []

    def limpa_tela(self):
        self.lb_selecionaArquivo["text"] = ""
        self.entryPalavrasChave.delete(0, END)
        self.lb_palavrasDefinidas["text"] = ""
        vetor_palavrasChave.clear()

    def abrir_janela_selecao_arquivo(self):
        self.caminho_arquivo = filedialog.askopenfilename(filetypes=[("Arquivos PDF", "*.pdf")])
        self.caminho.set(self.caminho_arquivo)

    def armazenar_palavrasChave(self):
        palavra_chave = self.entryPalavrasChave.get().lower()
        if palavra_chave != "":
            vetor_palavrasChave.append(palavra_chave)
            self.entryPalavrasChave.delete(0, END)
            self.atualizar_vetorPalavrasChave()

    def atualizar_vetorPalavrasChave(self):
        self.lb_palavrasDefinidas["text"] = ", ".join(vetor_palavrasChave)

    def extrair_texto_pdf(self):
        with open(self.caminho_arquivo, "rb") as arquivo_texto:
            reader = PdfReader(arquivo_texto)
            results = []

            for page in reader.pages:
                texto_pagina = page.extract_text()
                results.append(texto_pagina)
            return ' '.join(results).lower()

    def busca_texto_pdf(self):
        self.resultadoBusca.delete(1.0, END)

        self.textoCompleto = self.extrair_texto_pdf()

        self.nlp = spacy.load("pt_core_news_sm")

        self.doc = self.nlp(self.textoCompleto)

        self.palavras_chaves = vetor_palavrasChave

        self.encontrados = []

        for token in self.doc:
            if token.text in self.palavras_chaves:
                span = self.doc[token.i: token.i + 1].sent
                self.encontrados.append(span.text)

        if self.encontrados:
            for span in self.encontrados:
                self.resultadoBusca.insert(END, span + '\n')
            self.realcar_palavras_chave()
            self.salvar_arquivo()
        else:
            messagebox.showinfo("Sem retorno!", "Nenhuma palavra-chave pode ser encontrada!")

    def realcar_palavras_chave(self):
        texto = self.resultadoBusca.get("1.0", "end-1c")
        palavras_chave = vetor_palavrasChave

        for palavra_chave in palavras_chave:
            start_index = '1.0'
            while True:
                start_index = self.resultadoBusca.search(palavra_chave, start_index, stopindex="end-1c")
                if not start_index:
                    break
                end_index = f"{start_index}+{len(palavra_chave)}c"
                self.resultadoBusca.tag_add("bold", start_index, end_index)
                self.resultadoBusca.tag_config("bold", font=("Helvetica", 14, "bold"))
                start_index = end_index


class Application(Funcs):
    caminho = StringVar()

    def __init__(self):
        self.root = root
        self.tela()
        self.frames_da_tela()
        self.widgets_frame1()
        self.result_frame2()

        root.mainloop()

    def tela(self):
        self.root.title("Buscador de palavras PDF - Inovação MTEC")
        self.root.configure(background='#273f4b')
        self.root.geometry("640x480")
        self.root.resizable(True, True)
        self.root.maxsize(width=1920, height=1080)
        self.root.minsize(width=640, height=480)

    def frames_da_tela(self):
        self.frame_1 = Frame(self.root, bd=4, bg='#bbbbbb',
                             highlightbackground='#2e5169', highlightthickness=2.5)
        self.frame_1.place(relx=0.02, rely=0.02, relwidth=0.96, relheight=0.46)

        self.frame_2 = Frame(self.root, bd=4, bg='#bbbbbb',
                             highlightbackground='#2e5169', highlightthickness=2.5)
        self.frame_2.place(relx=0.02, rely=0.5, relwidth=0.96, relheight=0.46)

    def widgets_frame1(self):

        self.lb_apresentaArquivo = Label(self.frame_1, text="Arquivo PDF:", bg='#1d7881',
                                         fg='#d2dde0', font=('Helvetica', 11, 'bold'))
        self.lb_apresentaArquivo.place(relx=0.03, rely=0.1, relwidth=0.16, relheight=0.13)

        self.lb_selecionaArquivo = Label(self.frame_1, bg='#d2dde0',
                                         highlightbackground='#1d7881', highlightthickness=1.5,
                                         font=('Helvetica', 11), textvariable=self.caminho)
        self.lb_selecionaArquivo.place(relx=0.19, rely=0.1, relwidth=0.5, relheight=0.13)

        self.bt_addArquivo = Button(self.frame_1, text="+", bd=0.1, bg='#1d7881',
                                    fg='#d2dde0', font=('Helvetica', 10, 'bold'),
                                    command=self.abrir_janela_selecao_arquivo)
        self.bt_addArquivo.place(relx=0.68, rely=0.1, relwidth=0.04, relheight=0.13)

        self.lb_palavrasChave = Label(self.frame_1, text="Palavras:", bg='#1d7881',
                                      fg='#d2dde0', font=('Helvetica', 11, 'bold'))
        self.lb_palavrasChave.place(relx=0.03, rely=0.28, relwidth=0.16, relheight=0.13)

        self.lb_palavrasDefinidas = Label(self.frame_1, bg='#bbbbbb')
        self.lb_palavrasDefinidas.place(relx=0.03, rely=0.65, relwidth=0.88, relheight=0.13)

        self.entryPalavrasChave = Entry(self.frame_1, bg='#d2dde0',
                                        highlightbackground='#1d7881', highlightthickness=1.5,
                                        font=('Helvetica', 11))
        self.entryPalavrasChave.place(relx=0.19, rely=0.28, relwidth=0.5, relheight=0.13)

        self.bt_addPalavra = Button(self.frame_1, text="+", bd=0.1, bg='#1d7881',
                                    fg='#d2dde0', font=('Helvetica', 10, 'bold'),
                                    command=self.armazenar_palavrasChave)
        self.bt_addPalavra.place(relx=0.68, rely=0.28, relwidth=0.04, relheight=0.13)

        self.bt_buscar = Button(self.frame_1, text="Buscar", bd=2.5, bg='#1d7881',
                                fg='#d2dde0', font=('Helvetica', 10, 'bold'),
                                command=self.busca_texto_pdf)
        self.bt_buscar.place(relx=0.87, rely=0.1, relwidth=0.1, relheight=0.13)

        self.bt_limpar = Button(self.frame_1, text="Limpar", bd=2.5, bg='#1d7881',
                                fg='#d2dde0', font=('Helvetica', 10, 'bold'),
                                command=self.limpa_tela)
        self.bt_limpar.place(relx=0.87, rely=0.28, relwidth=0.1, relheight=0.13)

    def result_frame2(self):
        self.resultadoBusca = Text(self.frame_2, bg='#d2dde0', font=('Helvetica', 12))
        self.resultadoBusca.place(relx=0.01, rely=0.01, relwidth=0.95, relheight=0.95)

        self.scroolResultado = Scrollbar(self.frame_2, orient='vertical', command=self.resultadoBusca.yview)
        self.resultadoBusca.configure(yscrollcommand=self.scroolResultado.set)
        self.scroolResultado.place(relx=0.96, rely=0.01, relwidth=0.03, relheight=0.95)

    def abrir_janela_selecao_arquivo(self):
        self.caminho_arquivo = filedialog.askopenfilename(filetypes=[("Arquivos PDF", "*.pdf")])
        self.caminho.set(self.caminho_arquivo)

    def armazenar_palavrasChave(self):
        self.palavraChave = self.entryPalavrasChave.get().lower()
        if self.palavraChave != "":
            vetor_palavrasChave.append(self.palavraChave)
            self.entryPalavrasChave.delete(0, END)
            self.atualizar_vetorPalavrasChave()

    def atualizar_vetorPalavrasChave(self):
        self.lb_palavrasDefinidas["text"] = ", ".join(vetor_palavrasChave)

    def extrair_texto_pdf(self):
        with open(self.caminho.get(), "rb") as arquivo_texto:
            reader = PdfReader(arquivo_texto)
            results = []

            for i in range(0, len(reader.pages)):
                pagina_selecionada = reader.pages[i]
                texto_pagina = pagina_selecionada.extract_text()
                results.append(texto_pagina)
            return ' '.join(results).lower()

    def busca_texto_pdf(self):
        self.resultadoBusca.delete(1.0, END)

        self.textoCompleto = self.extrair_texto_pdf()

        self.nlp = spacy.load("pt_core_news_sm")

        self.doc = self.nlp(self.textoCompleto)

        self.palavras_chaves = vetor_palavrasChave

        self.encontrados = []

        for token in self.doc:
            if token.text in self.palavras_chaves:
                span = self.doc[token.i: token.i + 1].sent
                self.encontrados.append(span.text)

        if self.encontrados:
            for span in self.encontrados:
                self.resultadoBusca.insert(END, span + '\n')
            self.realcar_palavras_chave()
            self.salvar_arquivo()
        else:
            messagebox.showinfo("Sem retorno!", "Nenhuma palavra-chave pode ser encontrada!")

    def realcar_palavras_chave(self):
        texto = self.resultadoBusca.get("1.0", "end-1c")
        palavras_chave = vetor_palavrasChave

        for palavra_chave in palavras_chave:
            start_index = '1.0'
            while True:
                start_index = self.resultadoBusca.search(palavra_chave, start_index, stopindex="end-1c")
                if not start_index:
                    break
                end_index = f"{start_index}+{len(palavra_chave)}c"
                self.resultadoBusca.tag_add("bold", start_index, end_index)
                self.resultadoBusca.tag_config("bold", font=("Helvetica", 12, "bold"))
                start_index = end_index

    def salvar_arquivo(self):
        try:
            document = Document()
            document.add_heading('Palavras-Chave Encontradas:', 0)
            for span in self.encontrados:
                document.add_paragraph(span)

            caminho_arquivo = self.caminho.get().split('/')
            nome_arquivo = caminho_arquivo[-1].split('.pdf')[0]
            caminho_salvar = filedialog.asksaveasfilename(initialfile=nome_arquivo, defaultextension=".docx",
                                                          filetypes=[("Documento Word", "*.docx")])
            if caminho_salvar:
                document.save(caminho_salvar)
                messagebox.showinfo("Arquivo Salvo!", "O arquivo foi salvo com sucesso!")
        except Exception as e:
            messagebox.showinfo("Erro ao salvar arquivo", f"Ocorreu um erro ao salvar o arquivo: {str(e)}")


Application()
